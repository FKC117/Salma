"""
Report Generator for Word Document Creation

This service generates comprehensive Word documents from analysis results,
including charts, tables, and interpretations with customizable templates.
"""

import os
import logging
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from pathlib import Path
from django.conf import settings
from django.utils import timezone
from django.db import transaction
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from analytics.models import (
    ReportGeneration, AnalysisResult, GeneratedImage, User, 
    AnalysisSession, AuditTrail
)
from analytics.services.audit_trail_manager import AuditTrailManager

logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    Service for generating comprehensive Word reports
    """
    
    def __init__(self):
        self.audit_manager = AuditTrailManager()
        self.media_root = Path(settings.MEDIA_ROOT)
        self.reports_dir = self.media_root / 'reports'
        self.reports_dir.mkdir(parents=True, exist_ok=True)
        
        # Report settings
        self.default_template = 'analytical_report'
        self.max_images_per_report = 20
        self.max_tables_per_report = 50
    
    def generate_report(self, name: str, description: str, user: User,
                       analysis_results: List[AnalysisResult],
                       session: Optional[AnalysisSession] = None,
                       template_type: str = 'analytical_report',
                       custom_settings: Optional[Dict[str, Any]] = None) -> ReportGeneration:
        """
        Generate a comprehensive Word report
        
        Args:
            name: Report name
            description: Report description
            user: User generating report
            analysis_results: List of analysis results to include
            session: Analysis session
            template_type: Type of template to use
            custom_settings: Custom report settings
            
        Returns:
            ReportGeneration object
        """
        try:
            # Create report generation record
            with transaction.atomic():
                report = ReportGeneration.objects.create(
                    name=name,
                    description=description,
                    template_type=template_type,
                    custom_settings=custom_settings or {},
                    status='generating',
                    progress_percentage=0,
                    user=user,
                    session=session
                )
                
                # Add analysis results
                report.analysis_results.set(analysis_results)
            
            # Generate report content
            self._generate_report_content(report)
            
            # Update status
            report.status = 'completed'
            report.progress_percentage = 100
            report.completed_at = timezone.now()
            report.save()
            
            # Log audit trail
            self.audit_manager.log_user_action(
                user_id=user.id,
                action_type='generate_report',
                resource_type='report',
                resource_id=report.id,
                resource_name=name,
                action_description=f"Generated {template_type} report with {len(analysis_results)} analyses",
                success=True,
                data_changed=True
            )
            
            logger.info(f"Generated report {report.id} for user {user.id}")
            return report
            
        except Exception as e:
            logger.error(f"Report generation failed: {str(e)}", exc_info=True)
            
            # Update status to failed
            try:
                report.status = 'failed'
                report.error_message = str(e)
                report.save()
            except:
                pass
            
            raise ValueError(f"Report generation failed: {str(e)}")
    
    def _generate_report_content(self, report: ReportGeneration):
        """Generate the actual Word document content"""
        try:
            # Create new document
            doc = Document()
            
            # Set up document styles
            self._setup_document_styles(doc)
            
            # Add title page
            self._add_title_page(doc, report)
            
            # Add table of contents
            self._add_table_of_contents(doc, report)
            
            # Add executive summary
            self._add_executive_summary(doc, report)
            
            # Add analysis sections
            self._add_analysis_sections(doc, report)
            
            # Add conclusions and recommendations
            self._add_conclusions(doc, report)
            
            # Add appendices
            self._add_appendices(doc, report)
            
            # Save document
            self._save_report_document(report, doc)
            
        except Exception as e:
            logger.error(f"Report content generation failed: {str(e)}")
            raise
    
    def _setup_document_styles(self, doc: Document):
        """Set up document styles and formatting"""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Create custom styles
        self._create_custom_styles(doc)
    
    def _create_custom_styles(self, doc: Document):
        """Create custom styles for the report"""
        styles = doc.styles
        
        # Title style
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', 1)  # 1 = paragraph style
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Heading 1 style
        if 'Custom Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Custom Heading 1', 1)
            h1_style.font.name = 'Calibri'
            h1_style.font.size = Pt(18)
            h1_style.font.bold = True
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        if 'Custom Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Custom Heading 2', 1)
            h2_style.font.name = 'Calibri'
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.paragraph_format.space_before = Pt(10)
            h2_style.paragraph_format.space_after = Pt(4)
    
    def _add_title_page(self, doc: Document, report: ReportGeneration):
        """Add title page to document"""
        # Title
        title = doc.add_heading(report.name, 0)
        title.style = 'Custom Title'
        
        # Description
        if report.description:
            doc.add_paragraph(report.description)
        
        # Metadata
        doc.add_paragraph(f"Generated on: {timezone.now().strftime('%B %d, %Y at %I:%M %p')}")
        doc.add_paragraph(f"Generated by: {report.user.username}")
        
        if report.session:
            doc.add_paragraph(f"Analysis Session: {report.session.name}")
        
        # Page break
        doc.add_page_break()
    
    def _add_table_of_contents(self, doc: Document, report: ReportGeneration):
        """Add table of contents"""
        doc.add_heading('Table of Contents', level=1)
        
        toc_items = [
            "Executive Summary",
            "Analysis Results",
            "Conclusions and Recommendations",
            "Appendices"
        ]
        
        for item in toc_items:
            p = doc.add_paragraph()
            p.add_run(item).bold = True
            p.add_run(f" ......................... {len(toc_items) - toc_items.index(item)}")
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, report: ReportGeneration):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', level=1)
        
        # Summary content
        summary_text = f"""
        This report presents the results of {len(report.analysis_results.all())} analytical processes 
        conducted on the dataset. The analysis covers various statistical and data science techniques 
        to provide insights and recommendations.
        
        Key findings include:
        • Multiple analytical approaches were applied to the dataset
        • Statistical significance was assessed where applicable
        • Visualizations were generated to support findings
        • Recommendations are provided based on the analysis results
        """
        
        doc.add_paragraph(summary_text)
        doc.add_page_break()
    
    def _add_analysis_sections(self, doc: Document, report: ReportGeneration):
        """Add analysis results sections"""
        doc.add_heading('Analysis Results', level=1)
        
        for i, analysis in enumerate(report.analysis_results.all(), 1):
            # Analysis title
            doc.add_heading(f"Analysis {i}: {analysis.name}", level=2)
            
            # Analysis description
            if analysis.description:
                doc.add_paragraph(analysis.description)
            
            # Tool information
            if analysis.tool_used:
                doc.add_paragraph(f"Tool Used: {analysis.tool_used.display_name}")
            
            # Execution details
            doc.add_paragraph(f"Execution Time: {analysis.execution_time_ms}ms")
            doc.add_paragraph(f"Confidence Score: {analysis.confidence_score:.2f}")
            
            # Add content based on output type
            if analysis.output_type == 'table':
                self._add_table_content(doc, analysis)
            elif analysis.output_type == 'chart':
                self._add_chart_content(doc, analysis)
            elif analysis.output_type == 'text':
                self._add_text_content(doc, analysis)
            
            # Add spacing
            doc.add_paragraph()
    
    def _add_table_content(self, doc: Document, analysis: AnalysisResult):
        """Add table content to document"""
        try:
            data = analysis.result_data.get('data', [])
            if not data:
                return
            
            # Create table
            table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 0)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add headers
            if data:
                header_cells = table.rows[0].cells
                for i, header in enumerate(data[0]):
                    header_cells[i].text = str(header)
                    # Make header bold
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            
            # Add data rows
            for row_idx, row_data in enumerate(data[1:], 1):
                if row_idx < len(table.rows):
                    row_cells = table.rows[row_idx].cells
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(row_cells):
                            row_cells[col_idx].text = str(cell_data)
            
        except Exception as e:
            logger.error(f"Failed to add table content: {str(e)}")
            doc.add_paragraph("Error: Could not display table data")
    
    def _add_chart_content(self, doc: Document, analysis: AnalysisResult):
        """Add chart content to document"""
        try:
            # Get associated images
            images = analysis.generated_images.all()
            
            for image in images:
                if image.file_path:
                    # Add image to document
                    image_path = self.media_root / image.file_path
                    if image_path.exists():
                        doc.add_picture(str(image_path), width=Inches(6))
                        doc.add_paragraph(f"Figure: {image.name}")
            
            # Add chart description
            chart_data = analysis.result_data.get('chart_data', {})
            if chart_data.get('title'):
                doc.add_paragraph(f"Chart Title: {chart_data['title']}")
            
        except Exception as e:
            logger.error(f"Failed to add chart content: {str(e)}")
            doc.add_paragraph("Error: Could not display chart data")
    
    def _add_text_content(self, doc: Document, analysis: AnalysisResult):
        """Add text content to document"""
        try:
            text_content = analysis.result_data.get('text', '')
            if text_content:
                # Split into paragraphs
                paragraphs = text_content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
        except Exception as e:
            logger.error(f"Failed to add text content: {str(e)}")
            doc.add_paragraph("Error: Could not display text content")
    
    def _add_conclusions(self, doc: Document, report: ReportGeneration):
        """Add conclusions and recommendations section"""
        doc.add_heading('Conclusions and Recommendations', level=1)
        
        conclusions_text = """
        Based on the analysis results presented in this report, the following conclusions 
        and recommendations can be drawn:
        
        Conclusions:
        • The data analysis provides valuable insights into the dataset
        • Statistical patterns and trends have been identified
        • Visual representations support the analytical findings
        
        Recommendations:
        • Consider additional data collection if needed
        • Implement monitoring systems for ongoing analysis
        • Regular review of analytical results is recommended
        • Further investigation may be warranted in specific areas
        """
        
        doc.add_paragraph(conclusions_text)
    
    def _add_appendices(self, doc: Document, report: ReportGeneration):
        """Add appendices section"""
        doc.add_heading('Appendices', level=1)
        
        # Appendix A: Analysis Parameters
        doc.add_heading('Appendix A: Analysis Parameters', level=2)
        
        for i, analysis in enumerate(report.analysis_results.all(), 1):
            doc.add_paragraph(f"Analysis {i} Parameters:")
            if analysis.parameters_used:
                for key, value in analysis.parameters_used.items():
                    doc.add_paragraph(f"  {key}: {value}")
            doc.add_paragraph()
    
    def _save_report_document(self, report: ReportGeneration, doc: Document):
        """Save the Word document"""
        try:
            # Create user-specific directory
            user_dir = self.reports_dir / f"user_{report.user.id}"
            user_dir.mkdir(exist_ok=True)
            
            # Generate filename
            timestamp = int(timezone.now().timestamp())
            filename = f"report_{report.id}_{timestamp}.docx"
            file_path = user_dir / filename
            
            # Save document
            doc.save(str(file_path))
            
            # Update report record
            report.file_path = f"reports/user_{report.user.id}/{filename}"
            report.file_size_bytes = file_path.stat().st_size
            report.file_format = 'DOCX'
            report.generation_time_ms = int((timezone.now() - report.started_at).total_seconds() * 1000)
            report.save()
            
        except Exception as e:
            logger.error(f"Failed to save report document: {str(e)}")
            raise
    
    def get_report_download_url(self, report: ReportGeneration) -> str:
        """Get download URL for report"""
        return f"{settings.MEDIA_URL}{report.file_path}"
    
    def delete_report(self, report_id: int, user: User) -> bool:
        """Delete report and associated file"""
        try:
            report = ReportGeneration.objects.get(id=report_id, user=user)
            
            # Delete file
            if report.file_path:
                file_path = self.media_root / report.file_path
                if file_path.exists():
                    file_path.unlink()
            
            # Delete record
            report.delete()
            
            # Log audit trail
            self.audit_manager.log_user_action(
                user_id=user.id,
                action_type='delete_report',
                resource_type='report',
                resource_id=report_id,
                resource_name=report.name,
                action_description="Report deleted",
                success=True,
                data_changed=True
            )
            
            return True
            
        except ReportGeneration.DoesNotExist:
            return False
        except Exception as e:
            logger.error(f"Failed to delete report: {str(e)}")
            return False
    
    def list_user_reports(self, user: User, limit: int = 20) -> List[ReportGeneration]:
        """List reports for a user"""
        return ReportGeneration.objects.filter(user=user).order_by('-created_at')[:limit]